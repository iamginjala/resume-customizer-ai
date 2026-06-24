import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate(resume: dict) -> bytes:
    doc = Document()

    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # --- Helper Styling Functions ---
    def add_heading(text, size=11, bold=True):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = bold
        run.font.name = 'Helvetica'
        run.font.size = Pt(size)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        # --- XML Injection for Full-Width Bottom Border ---
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # Thickness of the line (6 eighths of a point)
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)
        # --------------------------------------------------
        
        return p

    def add_body(text, size=9.5, bold=False, italic=False, space_after=4):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Helvetica'
        run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        return p

    # --- Name + Contact Header ---
    name_p = doc.add_paragraph()
    name_run = name_p.add_run(resume.get("name", ""))
    name_run.bold = True
    name_run.font.name = 'Helvetica-Bold'
    name_run.font.size = Pt(18)
    name_p.paragraph_format.space_after = Pt(2)

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(resume.get("title", ""))
    title_run.bold = True
    title_run.font.name = 'Helvetica-Bold'
    title_run.font.size = Pt(11)
    title_p.paragraph_format.space_after = Pt(2)

    contact_parts = [resume.get("phone", ""), resume.get("email", ""), resume.get("linkedin", ""), resume.get("location", "")]
    contact_line = " | ".join(p for p in contact_parts if p)
    contact_p = doc.add_paragraph()
    contact_run = contact_p.add_run(contact_line)
    contact_run.font.name = 'Helvetica'
    contact_run.font.size = Pt(9)
    contact_p.paragraph_format.space_after = Pt(8)

    # --- Professional Summary ---
    if resume.get("summary"):
        summary_title = "PROFESSIONAL SUMMARY"
        add_heading(summary_title)
        add_body(resume.get("summary", ""))

    # --- Technical Skills ---
    skills = resume.get("skills", {})
    if skills:
        add_heading("Technical Skills")
        for category, items in skills.items():
            category_label = category.replace("_", " ").title()
            if "Platform" in category_label: category_label = "Power Platform"
            elif "Alm" in category_label: category_label = "Environment & ALM Administration"
            elif "Api" in category_label: category_label = "API & Data Integration"
            
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            
            label = p.add_run(category_label + ": ")
            label.bold = True
            label.font.name = 'Helvetica-Bold'
            label.font.size = Pt(9.5)
            
            skills_text = ", ".join(items) if isinstance(items, list) else items
            content = p.add_run(skills_text)
            content.font.name = 'Helvetica'
            content.font.size = Pt(9.5)

    # --- Experience ---
    experience = resume.get("experience", [])
    if experience:
        exp_title = "PROFESSIONAL EXPERIENCE" if any("project" in job or "client" in job for job in experience) else "EXPERIENCE"
        add_heading(exp_title)
        
        for job in experience:
            title = job.get("title", "")
            company = job.get("company", "")
            client_name = job.get("client", "")
            project = job.get("project", "")
            start = job.get("start_date", "")
            end = job.get("end_date", "")

            # Format primary company metadata heading
            company_line = f"{company} | {title}" if company and title else (company or title)
            add_body(company_line, size=10, bold=True, space_after=2)

            # Format secondary metadata (Client / Project) on a new line
            if client_name or project:
                extra_meta = []
                if client_name: extra_meta.append(f"Client: {client_name}")
                if project: extra_meta.append(f"Project: {project}")
                add_body(' | '.join(extra_meta), size=9.5, bold=False, space_after=2)

            date_line = f"{start} - {end}" if (start and end) else f"{start}{end}"
            if date_line:
                add_body(date_line, size=9.5, bold=True, space_after=4)

            context = job.get("context", "")
            if context:
                add_body(context)

            for bullet in job.get("bullets", []):
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(bullet)
                run.font.name = 'Helvetica'
                run.font.size = Pt(9.5)

    # --- Education ---
    education = resume.get("education", [])
    if education:
        add_heading("Education")
        for edu in education:
            degree = edu.get("degree", "")
            institution = edu.get("institution", "")
            start = edu.get("start_date", "")
            end = edu.get("end_date", "")
            gpa = edu.get("gpa", "")

            edu_line = f"{degree}, {institution}" if degree and institution else (degree or institution)
            add_body(edu_line, size=10, bold=True, space_after=2)

            edu_right = f"{start} - {end}" if (start and end) else f"{start}{end}"
            if gpa:
                edu_right += f" | GPA: {gpa}"
            if edu_right:
                add_body(edu_right, size=9.5, bold=True, space_after=4)

    # --- Certifications ---
    certifications = resume.get("certifications", [])
    if certifications:
        add_heading("Certifications")
        for cert in certifications:
            if isinstance(cert, dict):
                name = cert.get("name", "")
                issuer = cert.get("issuer", "")
                date = cert.get("date", "")
                cert_line = name
                if issuer: cert_line += f" | {issuer}"
                if date: cert_line += f" | {date}"
            else:
                cert_line = str(cert)
                
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(cert_line)
            run.font.name = 'Helvetica'
            run.font.size = Pt(9.5)

    # --- Projects ---
    projects = resume.get("projects", [])
    if projects:
        add_heading("Projects")
        for project in projects:
            name = project.get("name", "")
            description = project.get("description", "")
            tech = project.get("tech_stack", [])
            link = project.get("link", "")

            add_body(name, size=10, bold=True, space_after=2)
            if description:
                add_body(description, space_after=3)
            if tech:
                add_body(f"Tech: {', '.join(tech)}" if isinstance(tech, list) else f"Tech: {tech}", space_after=3)
            if link:
                add_body(f"Link: {link}", space_after=4)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()